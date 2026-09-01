from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
GAME_ROOT = ROOT / "Gamer"
CHARACTER_OUT = GAME_ROOT / "data" / "content" / "characters.json"
GACHA_POOL_OUT = GAME_ROOT / "data" / "content" / "gacha_pools.json"
UP_POOL_OUT = GAME_ROOT / "data" / "up_pool_config.json"
WEB_CHARACTER_OUT = GAME_ROOT / "web" / "src" / "data" / "characterPool.ts"

ATTRIBUTE_BY_INDEX = {
    1: ("WATER", "水"),
    2: ("EARTH", "土"),
    3: ("THUNDER", "雷"),
    4: ("WIND", "风"),
    5: ("FIRE", "火"),
    6: ("WOOD", "木"),
    7: ("LIGHT", "光"),
    8: ("DARK", "暗"),
}

PROFESSION_BY_SLOT = [
    ("PHYSICAL_TANK", "物理坦克"),
    ("MAGIC_TANK", "法系坦克"),
    ("PHYSICAL_MELEE_DPS", "物理近战输出"),
    ("MAGIC_MELEE_DPS", "法系近战输出"),
    ("PHYSICAL_RANGED_DPS", "物理远程输出"),
    ("MAGIC_RANGED_DPS", "法系远程输出"),
    ("HEALER", "治疗"),
    ("SUPPORT", "辅助"),
]


def read_docx_paragraphs(path: Path) -> list[str]:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return paragraphs


def after_colon(text: str) -> str:
    return re.sub(r"^.*?[：:]\s*", "", text).strip()


def normalize_name(name: str) -> str:
    name = name.strip().splitlines()[0].replace("代号：", "")
    name = re.split(r"\s+属性[：:]", name, maxsplit=1)[0].strip()
    name = re.sub(r"[（(].*?[）)]", "", name).strip()
    return name


def fallback_from_path(path: Path) -> tuple[str, str]:
    folder_name = path.parent.name
    filename = path.stem
    raw_name = filename if not filename.startswith("新建") else folder_name.split("-", 1)[-1]
    if "代号：" in raw_name:
        raw_name = raw_name.split("代号：", 1)[-1]
    profession_hint = re.sub(r"^\d+", "", folder_name.split("-", 1)[0]).strip()
    return normalize_name(raw_name), profession_hint


def parse_meta(
    paragraphs: list[str],
    fallback_name: str,
    fallback_profession: str,
    fallback_attribute: str,
) -> tuple[str, str, str, str, list[str]]:
    name = fallback_name
    attribute_name = fallback_attribute
    profession_name = fallback_profession
    weapon_name = ""
    description_parts: list[str] = []

    for text in paragraphs[:18]:
        if "角色名称" in text:
            name = normalize_name(after_colon(text))
        elif text.startswith("角色设定") and "：" in text:
            name = normalize_name(after_colon(text))
        elif text.startswith("属性") or "属性：" in text:
            match = re.search(r"[水土雷风火木光暗]", after_colon(text))
            if match:
                attribute_name = match.group(0)
        elif text.startswith("定位") or "定位：" in text:
            profession_name = normalize_profession_name(after_colon(text))
        elif text.startswith(("武器", "专属装备", "功法")) or "武器：" in text:
            weapon_name = after_colon(text)
        elif text.startswith(("特殊机制", "设计细节", "所属体系", "所属城池")):
            description_parts.append(text)

    return name, attribute_name, profession_name, weapon_name, description_parts


def normalize_profession_name(text: str) -> str:
    text = text.replace("法术", "法系")
    if "法系坦克" in text:
        return "法系坦克"
    if "物理坦克" in text:
        return "物理坦克"
    if "法系近战" in text:
        return "法系近战输出"
    if "物理近战" in text:
        return "物理近战输出"
    if "法系远程" in text:
        return "法系远程输出"
    if "物理远程" in text:
        return "物理远程输出"
    if "治疗" in text:
        return "治疗"
    if "辅助" in text:
        return "辅助"
    if "输出" in text:
        return "输出"
    if "坦克" in text:
        return "坦克"
    return text.strip()


def profession_key_from_name(name: str, slot_index: int) -> str:
    normalized = normalize_profession_name(name)
    for key, chinese_name in PROFESSION_BY_SLOT:
        if normalized == chinese_name:
            return key
    if normalized == "坦克":
        return PROFESSION_BY_SLOT[slot_index][0]
    if normalized == "输出":
        return PROFESSION_BY_SLOT[slot_index][0]
    return PROFESSION_BY_SLOT[slot_index][0]


def extract_skills(paragraphs: list[str]) -> list[dict[str, Any]]:
    skills: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    skill_re = re.compile(r"^(?:【[^】]+】\s*)?技能\s*([123])\s*[：:]\s*「?([^」\n]*)」?")

    for text in paragraphs:
        match = skill_re.match(text)
        if match:
            if current:
                skills.append(current)
            slot = int(match.group(1))
            skill_name = match.group(2).strip("「」：: ")
            current = {"slot": slot, "name": skill_name or f"技能{slot}", "effect": "", "notes": []}
            rest = text[match.end() :].strip()
            if rest:
                current["notes"].append(rest)
            continue

        if not current:
            continue
        if text.startswith("效果"):
            effect = after_colon(text)
            current["effect"] = f"{current['effect']}\n{effect}".strip() if current["effect"] else effect
        elif text.startswith(("释放细节", "技能台词", "机制说明", "备注", "说明")):
            current["notes"].append(text)
        elif re.match(r"^[三四五六七八九十]、", text):
            continue
        elif any(key in text for key in ["若", "当", "持续", "每层", "该技能", "20人团队", "复活", "护盾", "印记"]):
            current["notes"].append(text)

    if current:
        skills.append(current)

    by_slot = {skill["slot"]: skill for skill in skills}
    return [
        by_slot.get(slot, {"slot": slot, "name": f"技能{slot}", "effect": "", "notes": []})
        for slot in (1, 2, 3)
    ]


def make_character_payload() -> dict[str, Any]:
    characters: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []

    attribute_dirs = sorted(
        [path for path in ROOT.iterdir() if path.is_dir() and re.match(r"^[1-8]", path.name)],
        key=lambda path: int(path.name[:1]),
    )
    for attribute_dir in attribute_dirs:
        attribute_index = int(attribute_dir.name[:1])
        attribute_key, fallback_attribute_name = ATTRIBUTE_BY_INDEX[attribute_index]
        docx_files = sorted([path for path in attribute_dir.rglob("*.docx") if not path.name.startswith("~$")])

        for slot_index, docx_path in enumerate(docx_files):
            paragraphs = read_docx_paragraphs(docx_path)
            fallback_name, fallback_profession = fallback_from_path(docx_path)
            name, _attribute_name_from_doc, role_hint, weapon_name, description_parts = parse_meta(
                paragraphs,
                fallback_name,
                fallback_profession,
                fallback_attribute_name,
            )
            attribute_name = fallback_attribute_name
            profession_key, profession_name = PROFESSION_BY_SLOT[slot_index]
            number = (attribute_index - 1) * 8 + slot_index + 1
            skills = extract_skills(paragraphs)
            missing_slots = [skill["slot"] for skill in skills if not skill["effect"]]
            if missing_slots:
                warnings.append(
                    {
                        "name": name,
                        "source": docx_path.relative_to(ROOT).as_posix(),
                        "missingSkillEffectSlots": missing_slots,
                    }
                )

            overview = "\n".join(
                text
                for text in paragraphs
                if text.startswith(("定位", "武器", "专属装备", "特殊机制", "所属体系", "所属城池"))
            )
            characters.append(
                {
                    "id": f"char_{number:03d}_{attribute_key.lower()}_{profession_key.lower()}",
                    "number": number,
                    "name": name,
                    "attributeType": attribute_key,
                    "attributeName": attribute_name,
                    "professionType": profession_key,
                    "professionName": dict(PROFESSION_BY_SLOT).get(profession_key, profession_name),
                    "roleHint": role_hint,
                    "weaponName": weapon_name,
                    "rarity": "epic" if attribute_key in {"LIGHT", "DARK"} else "rare",
                    "source": docx_path.relative_to(ROOT).as_posix(),
                    "description": "\n".join(description_parts).strip() or overview,
                    "skills": skills,
                }
            )

    return {
        "version": 1,
        "source": "角色 Word 文档生成",
        "total": len(characters),
        "attributes": [
            {
                "key": attribute_key,
                "name": attribute_name,
                "count": sum(1 for character in characters if character["attributeType"] == attribute_key),
            }
            for attribute_key, attribute_name in ATTRIBUTE_BY_INDEX.values()
        ],
        "characters": characters,
        "generationWarnings": warnings,
    }


def write_outputs(payload: dict[str, Any]) -> None:
    CHARACTER_OUT.parent.mkdir(parents=True, exist_ok=True)
    CHARACTER_OUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    GACHA_POOL_OUT.write_text(
        json.dumps(make_gacha_pool_payload(payload["characters"]), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    UP_POOL_OUT.write_text(
        json.dumps(make_up_pool_config(payload["characters"]), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    WEB_CHARACTER_OUT.parent.mkdir(parents=True, exist_ok=True)
    WEB_CHARACTER_OUT.write_text(
        "export const CHARACTER_POOL = "
        + json.dumps(payload["characters"], ensure_ascii=False, indent=2)
        + " as const\n\nexport type CharacterPoolEntry = (typeof CHARACTER_POOL)[number]\n",
        encoding="utf-8",
    )


def to_gacha_entry(character: dict[str, Any], pool_key: str, weight: int = 100) -> dict[str, Any]:
    return {
        "entryId": f"{pool_key}_{character['id']}",
        "type": "character",
        "weight": weight,
        "rarity": character["rarity"],
        "characterConfigId": character["id"],
        "name": character["name"],
        "attributeType": character["attributeType"],
        "professionType": character["professionType"],
    }


def make_pool(key: str, name: str, description: str, characters: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "key": key,
        "name": name,
        "description": description,
        "cost": {"currency": "premiumCurrency", "amount": 160},
        "entries": [to_gacha_entry(character, key.lower()) for character in characters],
    }


def make_gacha_pool_payload(characters: list[dict[str, Any]]) -> dict[str, Any]:
    by_attribute = {
        attribute_key: [character for character in characters if character["attributeType"] == attribute_key]
        for attribute_key, _ in ATTRIBUTE_BY_INDEX.values()
    }
    up_ids = {"char_004_water_magic_melee_dps", "char_035_fire_physical_melee_dps", "char_052_light_magic_melee_dps"}
    up_entries = [
        to_gacha_entry(character, "up", 900 if character["id"] in up_ids else 100)
        for character in characters
    ]
    return {
        "version": 2,
        "description": "64角色配置抽卡池。旧Python前端使用大写key；Nest新后端也读取同一份配置。",
        "pools": [
            make_pool(
                "starter",
                "全属性角色池",
                "包含当前64名角色，适合静态试玩和新后端默认抽取。",
                characters,
            ),
            make_pool(
                "WATER_EARTH_THUNDER",
                "水土雷角色池",
                "包含水、土、雷三个属性的24名角色。",
                by_attribute["WATER"] + by_attribute["EARTH"] + by_attribute["THUNDER"],
            ),
            make_pool(
                "FIRE_WOOD_WIND",
                "火木风角色池",
                "包含火、木、风三个属性的24名角色。",
                by_attribute["FIRE"] + by_attribute["WOOD"] + by_attribute["WIND"],
            ),
            make_pool(
                "LIGHT_DARK",
                "光暗角色池",
                "包含光、暗两个属性的16名角色。",
                by_attribute["LIGHT"] + by_attribute["DARK"],
            ),
            {
                "key": "UP_POOL",
                "name": "当期UP角色池",
                "description": "UP角色权重更高，未命中时从64名角色中抽取。",
                "cost": {"currency": "premiumCurrency", "amount": 160},
                "entries": up_entries,
            },
        ],
    }


def make_up_pool_config(characters: list[dict[str, Any]]) -> dict[str, Any]:
    up_ids = {"char_004_water_magic_melee_dps", "char_035_fire_physical_melee_dps", "char_052_light_magic_melee_dps"}
    up_characters = [character for character in characters if character["id"] in up_ids]
    return {
        "pool_type": "UP_POOL",
        "up_rate": 0.5,
        "up_character_names": [character["name"] for character in up_characters],
        "title": "当期UP角色池",
        "description": "UP角色有50%概率优先抽取；未命中UP时从当前64名角色中抽取。",
    }


def main() -> None:
    payload = make_character_payload()
    write_outputs(payload)
    print(f"characters: {payload['total']}")
    print(f"warnings: {len(payload['generationWarnings'])}")
    if payload["generationWarnings"]:
        print(json.dumps(payload["generationWarnings"], ensure_ascii=False, indent=2))
    print(f"wrote: {CHARACTER_OUT}")
    print(f"wrote: {GACHA_POOL_OUT}")
    print(f"wrote: {UP_POOL_OUT}")
    print(f"wrote: {WEB_CHARACTER_OUT}")


if __name__ == "__main__":
    main()
